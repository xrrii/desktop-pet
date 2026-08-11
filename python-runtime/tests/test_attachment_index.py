from __future__ import annotations

import asyncio
from pathlib import Path

import openpyxl
import pytest
from docx import Document

from petdock_runtime.attachments.analysis import AttachmentAnalysisService
from petdock_runtime.attachments.index_store import AttachmentIndexStore
from petdock_runtime.attachments.store import AttachmentStore
from petdock_runtime.providers.embeddings import LocalHashEmbedding

"""C5 会话资料集直接注入、临时索引、来源位置和清理测试。"""


def _managed_item(root: Path, attachment_id: str, name: str, content: bytes) -> dict[str, object]:
    """在测试受控目录创建附件并返回登记参数。"""
    directory = root / attachment_id
    directory.mkdir(parents=True)
    path = directory / name
    path.write_bytes(content)
    return {
        "id": attachment_id,
        "name": name,
        "relativePath": str(path.relative_to(root)),
        "sizeBytes": path.stat().st_size,
    }


def _managed_existing(root: Path, attachment_id: str, path: Path) -> dict[str, object]:
    """返回已经由格式库写入受控目录的附件登记参数。"""
    return {
        "id": attachment_id,
        "name": path.name,
        "relativePath": str(path.relative_to(root)),
        "sizeBytes": path.stat().st_size,
    }


def _write_pdf(path: Path, text: str) -> None:
    """生成带单页 ASCII 文本层的最小 PDF。"""
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, value in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii") + value + b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    output.extend(b"".join(f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets[1:]))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(output)


def test_small_multiformat_dataset_is_injected_with_structured_locations(tmp_path: Path) -> None:
    """文本、CSV、PDF、DOCX、XLSX 小资料集应完整注入并保留结构位置。"""
    root = tmp_path / "attachments"
    items = [
        _managed_item(root, "1" * 32, "notes.txt", "文本口令 ALPHA。".encode()),
        _managed_item(root, "2" * 32, "prices.csv", "name,price\nA,10\n".encode()),
    ]

    pdf_dir = root / ("3" * 32)
    pdf_dir.mkdir(parents=True)
    pdf_path = pdf_dir / "manual.pdf"
    _write_pdf(pdf_path, "PDF ALPHA page")
    items.append(_managed_existing(root, "3" * 32, pdf_path))

    docx_dir = root / ("4" * 32)
    docx_dir.mkdir(parents=True)
    docx_path = docx_dir / "plan.docx"
    document = Document()
    document.add_heading("计划", level=1)
    document.add_paragraph("DOCX ALPHA 内容")
    document.save(docx_path)
    items.append(_managed_existing(root, "4" * 32, docx_path))

    xlsx_dir = root / ("5" * 32)
    xlsx_dir.mkdir(parents=True)
    xlsx_path = xlsx_dir / "data.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "数据"
    sheet["A1"] = "ALPHA"
    sheet["B1"] = 20
    workbook.save(xlsx_path)
    workbook.close()
    items.append(_managed_existing(root, "5" * 32, xlsx_path))

    store = AttachmentStore(str(tmp_path / "assistant.db"), str(root))
    embedding = LocalHashEmbedding()
    index = AttachmentIndexStore(":memory:", embedding)
    try:
        summaries = store.register_many(items)
        assert all(item["status"] == "ready" for item in summaries)
        store.bind_for_request([str(item["id"]) for item in items], "conversation-formats")
        service = AttachmentAnalysisService(store, index, embedding)

        context = asyncio.run(service.build_context("conversation-formats", "比较这些文件中的 ALPHA"))

        assert context.mode == "direct"
        assert context.total_attachments == 5
        assert len(context.sources) == 5
        assert "PDF ALPHA page" in context.context_text
        assert "DOCX ALPHA 内容" in context.context_text
        assert "工作表 数据" in context.context_text
        assert any(source["location"]["page"] == 1 for source in context.sources)
        assert any(source["location"]["sheet"] == "数据" for source in context.sources)
    finally:
        index.close()
        store.close()


def test_large_dataset_uses_isolated_index_and_covers_each_file(tmp_path: Path) -> None:
    """超过 Token 阈值后应使用独立索引，并为比较问题保留逐文件命中。"""
    root = tmp_path / "attachments"
    first_text = "\n".join(f"第 {index} 行 alpha timeout 配置为 30" for index in range(1, 180))
    second_text = "\n".join(f"第 {index} 行 alpha timeout 配置为 60" for index in range(1, 180))
    items = [
        _managed_item(root, "a" * 32, "first.txt", first_text.encode()),
        _managed_item(root, "b" * 32, "second.txt", second_text.encode()),
    ]
    store = AttachmentStore(str(tmp_path / "assistant.db"), str(root))
    embedding = LocalHashEmbedding()
    index_root = tmp_path / "session-index"
    index = AttachmentIndexStore(str(index_root), embedding)
    try:
        store.register_many(items)
        store.bind_for_request([str(item["id"]) for item in items], "conversation-large")
        service = AttachmentAnalysisService(store, index, embedding, direct_token_budget=20)

        context = asyncio.run(
            service.build_context(
                "conversation-large",
                "比较这些文件的 alpha timeout 配置",
            )
        )

        assert context.mode == "retrieval"
        assert {source["name"] for source in context.sources} == {"first.txt", "second.txt"}
        assert not context.unmatched_attachments
        assert all(source["mode"] == "retrieval" for source in context.sources)
        assert all(source["location"]["lineStart"] >= 1 for source in context.sources)
        assert (index_root / embedding.descriptor.signature / "index.db").is_file()
        assert (index_root / embedding.descriptor.signature / "chroma").is_dir()

        records = store.conversation_records("conversation-large")
        assert index.delete_conversation("conversation-large")
        assert index.search("conversation-large", records, "alpha timeout") == []
    finally:
        index.close()
        store.close()


def test_conversation_dataset_remains_available_without_resending_attachment_ids(tmp_path: Path) -> None:
    """附件绑定后，后续消息不重复提交附件 ID 也应继续使用会话资料集。"""
    root = tmp_path / "attachments"
    item = _managed_item(root, "c" * 32, "history.md", "会话持续口令：BLUE-42。".encode())
    store = AttachmentStore(str(tmp_path / "assistant.db"), str(root))
    embedding = LocalHashEmbedding()
    index = AttachmentIndexStore(":memory:", embedding)
    try:
        store.register_many([item])
        store.bind_for_request([str(item["id"])], "conversation-history")
        service = AttachmentAnalysisService(store, index, embedding)

        context = asyncio.run(service.build_context("conversation-history", "口令是什么"))

        assert context.mode == "direct"
        assert "BLUE-42" in context.context_text
        assert context.sources[0]["name"] == "history.md"
    finally:
        index.close()
        store.close()


def test_reconcile_removes_index_when_conversation_has_no_ready_attachment(tmp_path: Path) -> None:
    """启动核对时，只有失败附件的会话不得保留旧的派生索引。"""
    root = tmp_path / "attachments"
    ready = _managed_item(root, "e" * 32, "ready.txt", "alpha ready".encode())
    failed = _managed_item(root, "f" * 32, "failed.txt", b"\xff\xfe")
    store = AttachmentStore(str(tmp_path / "assistant.db"), str(root))
    embedding = LocalHashEmbedding()
    index = AttachmentIndexStore(":memory:", embedding)
    try:
        store.register_many([ready, failed])
        store.bind_for_request([str(ready["id"])], "conversation-reconcile")
        records = store.conversation_records("conversation-reconcile")
        index.sync_conversation("conversation-reconcile", records)
        assert "conversation-reconcile" in store.conversation_ids()
        store.delete_conversation("conversation-reconcile")

        index.reconcile(store.conversation_ids())

        assert not store.conversation_ids()
        assert not store.conversation_records("conversation-reconcile")
        assert index.search("conversation-reconcile", records, "alpha") == []
    finally:
        index.close()
        store.close()


def test_temporary_index_failure_reports_missing_coverage(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """临时索引故障时应保留对话并明确报告未读取范围。"""
    root = tmp_path / "attachments"
    content = "\n".join(f"第 {index} 行需要索引" for index in range(100))
    item = _managed_item(root, "d" * 32, "large.txt", content.encode())
    store = AttachmentStore(str(tmp_path / "assistant.db"), str(root))
    embedding = LocalHashEmbedding()
    index = AttachmentIndexStore(":memory:", embedding)

    def fail_sync(_conversation_id: str, _records: object) -> None:
        """模拟临时索引存储不可用。"""
        raise OSError("simulated")

    try:
        store.register_many([item])
        store.bind_for_request([str(item["id"])], "conversation-failure")
        monkeypatch.setattr(index, "sync_conversation", fail_sync)
        service = AttachmentAnalysisService(store, index, embedding, direct_token_budget=1)

        context = asyncio.run(service.build_context("conversation-failure", "总结附件"))

        assert context.mode == "retrieval"
        assert not context.sources
        assert context.unmatched_attachments[0]["name"] == "large.txt"
        assert any(item["code"] == "attachment_index_unavailable" for item in context.warnings)
        assert "本轮没有读取大资料集正文" in context.context_text
    finally:
        index.close()
        store.close()
