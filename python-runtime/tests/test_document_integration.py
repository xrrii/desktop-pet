"""C4 Parser Registry 在附件与知识库入口的集成和迁移测试。"""

from __future__ import annotations

import sqlite3
from pathlib import Path

from docx import Document
from PIL import Image

from petdock_runtime.attachments.store import AttachmentStore
from petdock_runtime.documents.parser import DocumentParserRegistry
from petdock_runtime.knowledge.service import KnowledgeService
from petdock_runtime.knowledge.store import KnowledgeStore


def test_attachment_uses_registry_and_returns_structured_preview(tmp_path: Path) -> None:
    """附件登记应通过统一 Registry，并在预览中返回结构块位置。"""
    root = tmp_path / "attachments"
    controlled = root / ("a" * 32)
    controlled.mkdir(parents=True)
    path = controlled / "sample.docx"
    document = Document()
    document.add_heading("附件标题", level=1)
    document.add_paragraph("附件正文")
    document.save(path)
    store = AttachmentStore(str(tmp_path / "assistant.db"), str(root))
    try:
        assert isinstance(store.registry, DocumentParserRegistry)
        summary = store.register_many([{
            "id": "a" * 32,
            "name": "sample.docx",
            "relativePath": f"{'a' * 32}/sample.docx",
            "sizeBytes": path.stat().st_size,
        }])[0]
        assert summary["status"] == "ready"
        assert summary["parserId"] == "docx-ooxml-v1"
        assert summary["blocks"][0]["location"]["headingPath"] == ["附件标题"]
        preview = store.preview("a" * 32, None, 0, 65_536)
        assert "附件正文" in preview["content"]
        assert preview["blocks"] == summary["blocks"]
        context, sources = store.build_context(store.get_records(["a" * 32]))
        assert "[位置：附件标题，段落 1]" in context
        assert sources[0]["location"]["headingPath"] == ["附件标题"]
    finally:
        store.close()


def test_image_registration_only_validates_metadata_until_send(tmp_path: Path) -> None:
    """图片登记阶段只做本地安全校验，不生成 Vision 摘要。"""
    root = tmp_path / "attachments"
    controlled = root / ("c" * 32)
    controlled.mkdir(parents=True)
    path = controlled / "photo.png"
    Image.new("RGB", (32, 24), "white").save(path)
    store = AttachmentStore(str(tmp_path / "assistant.db"), str(root))
    try:
        summary = store.register_many([{
            "id": "c" * 32,
            "name": "photo.png",
            "relativePath": f"{'c' * 32}/photo.png",
            "sizeBytes": path.stat().st_size,
        }])[0]
        assert summary["status"] == "ready"
        assert summary["parserId"] == "image-metadata-v1"
        assert summary["error"] is None
        preview = store.preview("c" * 32, None, 0, 1_000)
        assert "图片尺寸" in preview["content"]
        assert preview["blocks"][0]["kind"] == "image_metadata"
    finally:
        store.close()


def test_attachment_v1_database_migrates_without_losing_text(tmp_path: Path) -> None:
    """C1 旧附件表应增量增加结构列并继续读取既有正文。"""
    db_path = tmp_path / "legacy.db"
    connection = sqlite3.connect(db_path)
    connection.executescript(
        """
        CREATE TABLE attachments (
            id TEXT PRIMARY KEY, conversation_id TEXT, display_name TEXT NOT NULL,
            relative_storage_path TEXT NOT NULL, extension TEXT NOT NULL,
            detected_mime TEXT NOT NULL, size_bytes INTEGER NOT NULL, sha256 TEXT NOT NULL,
            parser_id TEXT, status TEXT NOT NULL, warning TEXT, error TEXT,
            text_content TEXT NOT NULL DEFAULT '', created_at TEXT NOT NULL, sent_at TEXT
        );
        CREATE TABLE attachment_schema_meta (key TEXT PRIMARY KEY, value TEXT NOT NULL);
        INSERT INTO attachments VALUES (
            'bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb', 'conversation-1', 'old.txt', 'x/old.txt',
            '.txt', 'text/plain', 3, 'hash', 'utf8-text-v1', 'ready', NULL, NULL,
            '旧文', '2026-01-01', '2026-01-01'
        );
        """
    )
    connection.commit()
    connection.close()
    store = AttachmentStore(str(db_path), str(tmp_path / "attachments"))
    try:
        record = store.get_records(["b" * 32])[0]
        assert record.text_content == "旧文"
        assert record.blocks_json == "[]"
    finally:
        store.close()


def test_knowledge_chunk_persists_document_location(tmp_path: Path) -> None:
    """知识库 Chunk 应保存并返回 PDF/Office 可引用位置。"""
    store = KnowledgeStore(str(tmp_path / "knowledge.db"))
    try:
        root = tmp_path / "library"
        root.mkdir()
        library = store.create_library("资料", str(root))
        _, records, _ = store.replace_document(
            str(library["id"]),
            "manual.pdf",
            1,
            100,
            "content-hash",
            [("第一页内容", 5, {"kind": "pdf_page", "value": "page-1", "page": 1})],
            "v2",
            "手册",
        )
        loaded = store.chunk_by_ids([str(records[0]["id"])])[str(records[0]["id"])]
        assert loaded["location"]["page"] == 1
        assert loaded["title"] == "手册"
    finally:
        store.close()


def test_attachment_and_knowledge_share_registry_instance(tmp_path: Path) -> None:
    """Server 注入时附件和知识库必须共享同一个 Registry 实例。"""
    registry = DocumentParserRegistry()
    attachments = AttachmentStore(
        str(tmp_path / "assistant.db"),
        str(tmp_path / "attachments"),
        registry,
    )
    knowledge_store = KnowledgeStore(str(tmp_path / "knowledge.db"))
    knowledge = KnowledgeService(
        knowledge_store,
        object(),  # type: ignore[arg-type] - 本测试不执行向量操作。
        parser_registry=registry,
    )
    try:
        assert attachments.registry is registry
        assert knowledge.registry is registry
    finally:
        attachments.close()
        knowledge_store.close()
